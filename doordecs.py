import cv2
# l_img = cv2.imread("frame.png")
# l_img = cv2.resize(l_img, (500, 500))
# s_img = cv2.imread("Famous_Scientists/Albert_Einstein.jpg")

names = ['Bowen', 'Fatma', 'Tri', 'Ekansh', 'Varun', 'Robert', 'Asheton', 'Asier', 'Tyler', 'Matthew',
 'Paden', 'Gage', 'Joshua', 'Grant', 'Xinpei', 'Yang', 'Ewart', 'Jingyang', 'Zilong', 'Orestis', 'Chung Mo',
  'Ikhyun', 'Ms. Edna', 'Aryaan', 'Ms. Pam', 'Vivek', 'Karthik', 'Nathaniel', 'Jordan', 'Maxwell', 'Shiva', 'Jack',
   'Andrew', 'Calvin', 'Quenesha', 'Jerred', 'Hayden', 'Sruthi', 'Isabela', 'Connor', 'Josh', 'Chido', 'Anh-Tu', 'Lauryn', 'Rohan',
    'Naveen', 'David',  'Braxton']

# print(len(names))


import os
from time import sleep

directory = '/Users/davidokao/Documents/ssm/doordecs/Famous_Scientists'
for filename,name in zip(os.listdir(directory), names):
    if filename.endswith(".jpg"):
    	l_img = cv2.imread("frame.png")
    	l_img = cv2.resize(l_img, (500, 500))
    	s_img = cv2.imread("Famous_Scientists/"+filename)
    	print(filename)
    	s_img = cv2.resize(s_img, (302, 337))

    	l_img[50: 50+s_img.shape[0], 100: 100+s_img.shape[1]] = s_img

    	font = cv2.FONT_HERSHEY_COMPLEX
    	font2 = cv2.FONT_HERSHEY_COMPLEX_SMALL
    	text = name
    	text2 = filename[:-4].replace('_', ' ')
    	textsize = cv2.getTextSize(text, font, 1, 2)[0]
    	textsize2 = cv2.getTextSize(text2, font2, 1, 2)[0]
    	textX = (l_img.shape[1] - textsize[0]) // 2
    	textY = 480
    	textx = (l_img.shape[1] - textsize2[0]) // 2
    	texty = 43
    	cv2.putText(l_img, text, (textX, textY ), font, 1, (0, 0, 0), 2)
    	cv2.putText(l_img, text2, (textx, texty ), font2, 1, (0, 0, 0), 2)
    	cv2.imshow('tttttt',l_img)
    	cv2.imwrite(name+'.jpg', l_img)
    	# cv2.waitKey(1000)


from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()
# r.add_text('Good Morning every body,This is my ')
for i in names:
	r.add_picture(i+'.jpg')
# r.add_text(' do you like it?')

document.save('demo.docx')